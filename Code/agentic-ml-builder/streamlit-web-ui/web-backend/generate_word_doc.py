import sys
import subprocess
import os
import re

def install(package):
    print(f"Installing {package}...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    print("python-docx library not found. Installing...")
    install("python-docx")
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_doc(md_file, docx_file):
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found.")
        return

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Title Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Heading 1
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            # Heading 2
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Heading 3
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('* ') or line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
             # Numbered list (simple approximation)
            p = doc.add_paragraph(line, style='List Number')
        else:
            # Normal text
            p = doc.add_paragraph(line)

    try:
        doc.save(docx_file)
        print(f"Successfully created {docx_file}")
    except PermissionError:
        new_file = docx_file.replace(".docx", "_v2.docx")
        print(f"Warning: Could not write to {docx_file} (file might be open).")
        print(f"Saving to {new_file} instead...")
        doc.save(new_file)
        print(f"Successfully created {new_file}")

if __name__ == "__main__":
    create_word_doc("Project_Documentation.md", "Project_Documentation.docx")
